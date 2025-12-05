"""
Document Loader Node - Load and extract text from various document formats.
This node loads documents (PDF, DOCX, TXT, MD) and extracts text content.
"""

from typing import Dict, Any, List, Optional
import sys
import os

# Add the parent directory to the path to import base_node and ui_components
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from base_node import BaseNode, NodeInput, NodeOutput, NodeParameter, NodeStyling
from ui_components import (
    NodeUIConfig, UIGroup, DialogConfig,
    create_text_input, create_select, create_label, create_divider, create_file_upload,
    UIOption
)


class DocumentLoaderNode(BaseNode):
    """
    Document Loader Node - Extract text from documents.

    This node loads various document formats (PDF, DOCX, TXT, MD) and
    extracts their text content for use in workflows.
    """
    
    def _define_required_credentials(self, parameters: Optional[Dict[str, Any]] = None) -> List[str]:
        """No credentials required for DocumentLoaderNode"""
        return []

    def _define_category(self) -> str:
        """Define category for DocumentLoaderNode"""
        return "Data/Files"

    def _define_inputs(self) -> List[NodeInput]:
        """Define the input structure for DocumentLoaderNode"""
        return []  # No inputs - file is uploaded via parameter

    def _define_outputs(self) -> List[NodeOutput]:
        """Define the output structure for DocumentLoaderNode"""
        return [
            NodeOutput(
                name="text",
                type="string",
                description="Extracted text content from the document"
            ),
            NodeOutput(
                name="metadata",
                type="object",
                description="Document metadata (file name, type, size, page count)"
            )
        ]

    def _define_parameters(self) -> List[NodeParameter]:
        """Define the parameters for DocumentLoaderNode"""
        return [
            NodeParameter(
                name="uploaded_file",
                type="file",
                description="Upload a document file (PDF, DOCX, TXT, MD)",
                required=True,
                default_value=""
            ),
            NodeParameter(
                name="file_type",
                type="string",
                description="Document file type (auto-detect if empty)",
                required=False,
                default_value="auto",
                options=["auto", "pdf", "docx", "txt", "md"]
            ),
            NodeParameter(
                name="encoding",
                type="string",
                description="Text encoding for TXT/MD files",
                required=False,
                default_value="utf-8"
            )
        ]

    def _define_styling(self) -> NodeStyling:
        """Define custom styling for DocumentLoaderNode"""
        return NodeStyling(
            html_template="""
            <div class="doc-loader-node-container">
                <div class="doc-loader-icon">
                    <svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" class="lucide lucide-file-text"><path d="M15 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V7Z"/><path d="M14 2v4a2 2 0 0 0 2 2h4"/><path d="M10 9H8"/><path d="M16 13H8"/><path d="M16 17H8"/></svg>
                </div>
                <div class="doc-loader-content">
                    <div class="doc-loader-title">Load Document</div>
                    <div class="doc-loader-subtitle">READ FILE</div>
                </div>
            </div>
            """,
            custom_css="""
            .doc-loader-node-container {
                display: flex;
                align-items: center;
                padding: 16px 20px;
                background: #1f1f1f;
                border: 1.5px solid #10b981;
                border-radius: 12px;
                box-shadow: 0 2px 8px rgba(0, 0, 0, 0.3);
                transition: all 0.2s ease;
                transform-origin: center center;
                width: 220px;
                height: 100px;
                position: relative;
            }
            .doc-loader-node-container:hover {
                border-color: #34d399;
                box-shadow: 0 4px 12px rgba(16, 185, 129, 0.2);
            }
            .doc-loader-icon { margin-right: 12px; flex-shrink: 0; color: #10b981; display: flex; align-items: center; }
            .doc-loader-icon svg { width: 20px; height: 20px; }
            .doc-loader-content { flex: 1; display: flex; flex-direction: column; justify-content: center; gap: 2px; }
            .doc-loader-title { font-size: 13px; font-weight: 600; color: #ffffff; margin-bottom: 2px; line-height: 1.2; }
            .doc-loader-subtitle { font-size: 11px; color: #10b981; opacity: 0.9; line-height: 1.2; font-weight: 700; letter-spacing: 0.5px; text-transform: uppercase; }
            """,
            icon="<svg xmlns=\"http://www.w3.org/2000/svg\" width=\"24\" height=\"24\" viewBox=\"0 0 24 24\" fill=\"none\" stroke=\"currentColor\" stroke-width=\"2\" stroke-linecap=\"round\" stroke-linejoin=\"round\" class=\"lucide lucide-file-text\"><path d=\"M15 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V7Z\"/><path d=\"M14 2v4a2 2 0 0 0 2 2h4\"/><path d=\"M10 9H8\"/><path d=\"M16 13H8\"/><path d=\"M16 17H8\"/></svg>",
            subtitle="READ FILE",
            background_color="#1f1f1f",
            border_color="#10b981",
            text_color="#ffffff",
            shape="rounded",
            width=220,
            height=100,
            css_classes="",
            inline_styles='{}',
            icon_position=""
        )

    def _define_ui_config(self) -> NodeUIConfig:
        """Define the UI configuration for DocumentLoaderNode"""
        return NodeUIConfig(
            node_id=self.node_id,
            node_name="DocumentLoaderNode",
            groups=[
                UIGroup(
                    name="file_upload_config",
                    label="Upload Document",
                    components=[
                        create_label(
                            text="Upload a document to extract text content. Supports PDF, DOCX, TXT, and MD files."
                        ),
                        create_divider(),
                        create_file_upload(
                            name="uploaded_file",
                            label="Choose File *",
                            required=True,
                            accept=".pdf,.docx,.doc,.txt,.md"
                        )
                    ],
                    styling={
                        "background": "#2a2a2a",
                        "border_radius": "12px",
                    }
                ),
                UIGroup(
                    name="file_config",
                    label="File Options",
                    description="Configure how the file should be processed",
                    components=[
                        create_select(
                            name="file_type",
                            label="File Type",
                            description="Specify the file type manually or let the system auto-detect it",
                            required=False,
                            default_value="auto",
                            options=[
                                UIOption(value="auto", label="Auto-detect"),
                                UIOption(value="pdf", label="PDF (.pdf)"),
                                UIOption(value="docx", label="Word Document (.docx)"),
                                UIOption(value="txt", label="Text File (.txt)"),
                                UIOption(value="md", label="Markdown (.md)")
                            ],
                            searchable=False
                        ),
                        create_text_input(
                            name="encoding",
                            label="Text Encoding",
                            description="Character encoding for text files (usually utf-8)",
                            required=False,
                            default_value="utf-8",
                            placeholder="utf-8"
                        )
                    ],
                    styling={
                        "background": "#2a2a2a",
                        "border_radius": "12px",
                    }
                )
            ],
            layout="vertical",
            global_styling={
                "font_family": "Inter, sans-serif",
                "color_scheme": "light"
            },
            dialog_config=DialogConfig(
                title="Configure Document Loader",
                description="Load and extract text from documents (PDF, DOCX, TXT, MD)",
                background_color="#1f1f1f",
                border_color="#10b981",
                text_color="#ffffff",
                icon="""<svg xmlns="http://www.w3.org/2000/svg" width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" x2="8" y1="13" y2="13"/><line x1="16" x2="8" y1="17" y2="17"/><polyline points="10 9 9 9 8 9"/></svg>""",
                icon_color="#10b981",
                header_background="#1f1f1f",
                footer_background="#1f1f1f",
                button_primary_color="#10b981",
                button_secondary_color="#374151"
            )
        )

    def _extract_pdf(self, file_path: str) -> tuple[str, dict]:
        """Extract text from PDF file"""
        try:
            import pypdf
            text_content = []
            metadata = {}

            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                metadata['page_count'] = len(pdf_reader.pages)
                metadata['metadata'] = pdf_reader.metadata if pdf_reader.metadata else {}

                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())

            return '\n\n'.join(text_content), metadata
        except ImportError:
            return "", {"error": "pypdf library not installed. Install with: pip install pypdf"}
        except Exception as e:
            return "", {"error": f"Failed to extract PDF: {str(e)}"}

    def _extract_docx(self, file_path: str) -> tuple[str, dict]:
        """Extract text from DOCX file"""
        try:
            import docx
            doc = docx.Document(file_path)
            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'section_count': len(doc.sections)
            }

            return '\n\n'.join(text_content), metadata
        except ImportError:
            return "", {"error": "python-docx library not installed. Install with: pip install python-docx"}
        except Exception as e:
            return "", {"error": f"Failed to extract DOCX: {str(e)}"}

    def _extract_text_file(self, file_path: str, encoding: str = "utf-8") -> tuple[str, dict]:
        """Extract text from TXT or MD file"""
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()

            metadata = {
                'line_count': len(content.splitlines()),
                'char_count': len(content),
                'encoding': encoding
            }

            return content, metadata
        except Exception as e:
            return "", {"error": f"Failed to read text file: {str(e)}"}

    def execute(self, inputs: Dict[str, Any], parameters: Dict[str, Any]) -> Dict[str, Any]:
        """
        Execute the DocumentLoaderNode logic

        Args:
            inputs: Empty dictionary (no inputs)
            parameters: Dictionary containing uploaded_file, file_type, encoding settings

        Returns:
            Dictionary containing extracted text and metadata
        """
        # Handle file upload data structure
        uploaded_file = parameters.get("uploaded_file", "")

        file_path = ""

        # Extract file path from upload data
        if uploaded_file:
            # If it's a dict (file upload response from API), extract the path
            if isinstance(uploaded_file, dict):
                # Check if it's the response from /api/v1/files/upload
                if "file" in uploaded_file and isinstance(uploaded_file["file"], dict):
                    file_path = uploaded_file["file"].get("path", "")
                # Check for direct path
                elif "path" in uploaded_file:
                    file_path = uploaded_file.get("path", "")
                # Legacy support: nested structure like {'0': {...}}
                elif not file_path:
                    for key, value in uploaded_file.items():
                        if isinstance(value, dict):
                            file_path = (
                                value.get("path") or
                                value.get("filepath") or
                                ""
                            )
                            if file_path:
                                break
            # If it's a string, use it directly as path
            else:
                file_path = str(uploaded_file).strip() if uploaded_file else ""

        if not file_path:
            return {
                "text": "",
                "success": False,
                "metadata": {
                    "error": "No file uploaded. Please upload a document file."
                }
            }

        # Check if file exists
        if not os.path.exists(file_path):
            return {
                "text": "",
                "success": False,
                "metadata": {"error": f"File not found: {file_path}"}
            }

        # Get file info
        file_size = os.path.getsize(file_path)
        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_name)[1].lower()

        # Get parameters
        file_type = parameters.get("file_type", "auto")
        encoding = parameters.get("encoding", "utf-8")

        # Auto-detect file type
        if file_type == "auto":
            ext_map = {
                '.pdf': 'pdf',
                '.docx': 'docx',
                '.doc': 'docx',
                '.txt': 'txt',
                '.md': 'md',
                '.markdown': 'md'
            }
            file_type = ext_map.get(file_ext, 'txt')

        # Extract text based on file type
        text = ""
        metadata = {
            "file_name": file_name,
            "file_size": file_size,
            "file_type": file_type
        }

        if file_type == "pdf":
            text, extra_metadata = self._extract_pdf(file_path)
            metadata.update(extra_metadata)
        elif file_type == "docx":
            text, extra_metadata = self._extract_docx(file_path)
            metadata.update(extra_metadata)
        elif file_type in ["txt", "md"]:
            text, extra_metadata = self._extract_text_file(file_path, encoding)
            metadata.update(extra_metadata)
        else:
            return {
                "text": "",
                "success": False,
                "metadata": {
                    "error": f"Unsupported file type: {file_type}",
                    "file_name": file_name,
                    "file_size": file_size,
                    "file_type": file_type
                }
            }
        
        # Check if extraction resulted in an error
        if metadata.get("error"):
            return {
                "text": "",
                "success": False,
                "metadata": metadata
            }

        # Store result in node data for display
        self.node_data = {
            "file_name": file_name,
            "file_type": file_type
        }

        return {
            "text": text,
            "metadata": metadata
        }
